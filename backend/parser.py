import textract
import re
from datetime import datetime
import PyPDF2
import docx
from pathlib import Path

def parse_schedule(filepath):
    """
    Enhanced parser to extract dates, deadlines, and course information
    from various file formats (PDF, DOCX, TXT)
    """
    file_path = Path(filepath)
    extension = file_path.suffix.lower()
    
    try:
        if extension == '.pdf':
            text = extract_from_pdf(filepath)
        elif extension == '.docx':
            text = extract_from_docx(filepath)
        elif extension == '.txt':
            text = extract_from_txt(filepath)
        else:
            # Fallback to textract for other formats
            text = textract.process(filepath).decode('utf-8')
        
        # Parse the extracted text
        parsed_data = parse_text_content(text)
        return parsed_data
        
    except Exception as e:
        print(f"Error parsing file: {e}")
        return {'dates': [], 'courses': [], 'deadlines': []}

def extract_from_pdf(filepath):
    """Extract text from PDF files"""
    text = ""
    try:
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except:
        # Fallback to textract if PyPDF2 fails
        text = textract.process(filepath).decode('utf-8')
    
    return text

def extract_from_docx(filepath):
    """Extract text from DOCX files"""
    try:
        doc = docx.Document(filepath)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    except:
        # Fallback to textract
        return textract.process(filepath).decode('utf-8')

def extract_from_txt(filepath):
    """Extract text from TXT files"""
    with open(filepath, 'r', encoding='utf-8') as file:
        return file.read()

def parse_text_content(text):
    """
    Parse the extracted text to find dates, courses, and deadlines
    """
    parsed_data = {
        'dates': [],
        'courses': [],
        'deadlines': [],
        'class_times': []
    }
    
    # Clean up the text
    text = text.replace('\n', ' ').replace('\r', ' ')
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
    
    # Extract various date formats
    dates = extract_dates(text)
    parsed_data['dates'] = dates
    
    # Extract course information
    courses = extract_courses(text)
    parsed_data['courses'] = courses
    
    # Extract deadlines and assignments
    deadlines = extract_deadlines(text)
    parsed_data['deadlines'] = deadlines
    
    # Extract class times
    class_times = extract_class_times(text)
    parsed_data['class_times'] = class_times
    
    return parsed_data

def extract_dates(text):
    """Extract various date formats from text"""
    dates = []
    
    # Different date patterns
    date_patterns = [
        r'\d{1,2}/\d{1,2}/\d{4}',           # MM/DD/YYYY or M/D/YYYY
        r'\d{1,2}-\d{1,2}-\d{4}',           # MM-DD-YYYY
        r'\d{4}-\d{1,2}-\d{1,2}',           # YYYY-MM-DD
        r'\b\w+ \d{1,2}, \d{4}\b',          # Month DD, YYYY
        r'\b\d{1,2} \w+ \d{4}\b',           # DD Month YYYY
        r'\b\w+ \d{1,2}\b',                 # Month DD (current year assumed)
    ]
    
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        dates.extend(matches)
    
    # Remove duplicates and sort
    dates = list(set(dates))
    
    # Validate and format dates
    validated_dates = []
    for date_str in dates:
        validated_date = validate_and_format_date(date_str)
        if validated_date:
            validated_dates.append(validated_date)
    
    return validated_dates

def extract_courses(text):
    """Extract course codes and names"""
    courses = []
    
    # Course code patterns (e.g., CS101, MATH 201, ENGL-102)
    course_patterns = [
        r'\b[A-Z]{2,4}[- ]?\d{3,4}[A-Z]?\b',  # CS101, MATH 201, ENGL-102A
        r'\b[A-Z]{2,4} \d{3,4}\b',            # CS 101, MATH 201
    ]
    
    for pattern in course_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        courses.extend([match.upper() for match in matches])
    
    # Extract course names (lines that contain "course" or appear after course codes)
    course_name_pattern = r'(?:course|class):\s*(.+?)(?:\n|$)'
    course_names = re.findall(course_name_pattern, text, re.IGNORECASE)
    
    return {
        'codes': list(set(courses)),
        'names': course_names
    }

def extract_deadlines(text):
    """Extract deadlines, assignments, and exams"""
    deadlines = []
    
    # Keywords that indicate deadlines
    deadline_keywords = [
        'due', 'deadline', 'assignment', 'exam', 'test', 'quiz', 
        'project', 'paper', 'presentation', 'final', 'midterm'
    ]
    
    # Create pattern to find deadline-related sentences
    keyword_pattern = '|'.join(deadline_keywords)
    
    # Find sentences containing deadline keywords
    sentences = re.split(r'[.!?]', text)
    
    for sentence in sentences:
        if re.search(keyword_pattern, sentence, re.IGNORECASE):
            # Extract dates from this sentence
            sentence_dates = extract_dates(sentence)
            
            # Determine deadline type
            deadline_type = determine_deadline_type(sentence)
            
            # Extract title/description
            title = extract_deadline_title(sentence)
            
            if sentence_dates:
                for date in sentence_dates:
                    deadlines.append({
                        'date': date,
                        'type': deadline_type,
                        'title': title,
                        'description': sentence.strip(),
                        'source': 'parsed'
                    })
    
    return deadlines

def extract_class_times(text):
    """Extract class meeting times"""
    class_times = []
    
    # Time patterns
    time_patterns = [
        r'\d{1,2}:\d{2}\s*[AaPp][Mm]',     # 10:30 AM
        r'\d{1,2}:\d{2}',                   # 10:30 (24-hour)
        r'\d{1,2}\s*[AaPp][Mm]',           # 10 AM
    ]
    
    # Day patterns
    day_patterns = [
        r'\b(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)\b',
        r'\b(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)\b',
        r'\b(?:M|T|W|R|F|S|U)\b',  # Single letter abbreviations
        r'\b(?:MWF|TTH|TR|MW|WF)\b'  # Common combinations
    ]
    
    # Look for patterns like "MWF 10:00 AM - 11:00 AM"
    schedule_pattern = r'([MTWRFSU]+)\s+(\d{1,2}:\d{2}\s*[AaPp][Mm])\s*(?:-\s*(\d{1,2}:\d{2}\s*[AaPp][Mm]))?'
    
    matches = re.findall(schedule_pattern, text, re.IGNORECASE)
    
    for match in matches:
        days, start_time, end_time = match
        class_times.append({
            'days': days,
            'start_time': start_time,
            'end_time': end_time if end_time else None
        })
    
    return class_times

def determine_deadline_type(sentence):
    """Determine the type of deadline from the sentence"""
    sentence_lower = sentence.lower()
    
    if any(word in sentence_lower for word in ['exam', 'test', 'final', 'midterm']):
        return 'exam'
    elif any(word in sentence_lower for word in ['assignment', 'homework', 'hw']):
        return 'assignment'
    elif any(word in sentence_lower for word in ['project', 'paper']):
        return 'project'
    elif 'quiz' in sentence_lower:
        return 'quiz'
    elif any(word in sentence_lower for word in ['presentation', 'present']):
        return 'presentation'
    else:
        return 'assignment'  # Default

def extract_deadline_title(sentence):
    """Extract a meaningful title from the deadline sentence"""
    sentence = sentence.strip()
    
    # Remove common prefixes
    prefixes_to_remove = [
        r'^.*?due:?\s*',
        r'^.*?deadline:?\s*',
        r'^.*?assignment:?\s*',
        r'^.*?exam:?\s*',
        r'^.*?test:?\s*',
        r'^.*?quiz:?\s*',
        r'^.*?project:?\s*'
    ]
    
    title = sentence
    for prefix in prefixes_to_remove:
        title = re.sub(prefix, '', title, flags=re.IGNORECASE)
    
    # Clean up and limit length
    title = re.sub(r'\s+', ' ', title).strip()
    if len(title) > 100:
        title = title[:97] + "..."
    
    return title if title else "Untitled Deadline"

def validate_and_format_date(date_str):
    """Validate and format date strings to YYYY-MM-DD"""
    try:
        # Try different date formats
        date_formats = [
            '%m/%d/%Y',   # MM/DD/YYYY
            '%m-%d-%Y',   # MM-DD-YYYY
            '%Y-%m-%d',   # YYYY-MM-DD
            '%B %d, %Y',  # Month DD, YYYY
            '%d %B %Y',   # DD Month YYYY
            '%b %d, %Y',  # Mon DD, YYYY
            '%d %b %Y',   # DD Mon YYYY
            '%m/%d',      # MM/DD (current year)
            '%B %d',      # Month DD (current year)
            '%b %d'       # Mon DD (current year)
        ]
        
        for date_format in date_formats:
            try:
                if '%Y' not in date_format:
                    # Add current year for partial dates
                    current_year = datetime.now().year
                    date_str_with_year = f"{date_str}, {current_year}"
                    date_format_with_year = f"{date_format}, %Y"
                    parsed_date = datetime.strptime(date_str_with_year, date_format_with_year)
                else:
                    parsed_date = datetime.strptime(date_str, date_format)
                
                return parsed_date.strftime('%Y-%m-%d')
            except ValueError:
                continue
        
        return None
    except:
        return None

def extract_semester_info(text):
    """Extract semester/term information"""
    semester_patterns = [
        r'\b(Fall|Spring|Summer|Winter)\s+(\d{4})\b',
        r'\b(\d{4})\s+(Fall|Spring|Summer|Winter)\b',
        r'\bSemester:\s*(.+?)(?:\n|$)',
        r'\bTerm:\s*(.+?)(?:\n|$)'
    ]
    
    for pattern in semester_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0).strip()
    
    return None

def extract_instructor_info(text):
    """Extract instructor information"""
    instructor_patterns = [
        r'(?:Instructor|Professor|Dr\.|Prof\.)\s*:?\s*(.+?)(?:\n|$)',
        r'(?:Taught by|Teacher)\s*:?\s*(.+?)(?:\n|$)',
        r'Dr\.\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
        r'Prof\.\s+([A-Z][a-z]+\s+[A-Z][a-z]+)'
    ]
    
    for pattern in instructor_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return None

def smart_parse_syllabus(filepath):
    """
    Advanced syllabus parsing that attempts to understand document structure
    """
    parsed_data = parse_schedule(filepath)
    
    # Extract additional information
    try:
        if filepath.endswith('.pdf'):
            text = extract_from_pdf(filepath)
        elif filepath.endswith('.docx'):
            text = extract_from_docx(filepath)
        else:
            text = extract_from_txt(filepath)
        
        # Add additional parsing
        parsed_data['semester'] = extract_semester_info(text)
        parsed_data['instructor'] = extract_instructor_info(text)
        
        # Try to identify syllabus sections
        sections = identify_syllabus_sections(text)
        parsed_data['sections'] = sections
        
    except Exception as e:
        print(f"Error in smart parsing: {e}")
    
    return parsed_data

def identify_syllabus_sections(text):
    """Identify common syllabus sections"""
    sections = {}
    
    section_patterns = {
        'grading': r'(?:grading|grade breakdown|assessment)[\s\S]*?(?=\n\n|\n[A-Z])',
        'attendance': r'(?:attendance|absence)[\s\S]*?(?=\n\n|\n[A-Z])',
        'schedule': r'(?:schedule|calendar|timeline)[\s\S]*?(?=\n\n|\n[A-Z])',
        'policies': r'(?:policies|policy|rules)[\s\S]*?(?=\n\n|\n[A-Z])',
        'objectives': r'(?:objectives|goals|outcomes)[\s\S]*?(?=\n\n|\n[A-Z])'
    }
    
    for section_name, pattern in section_patterns.items():
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            sections[section_name] = match.group(0).strip()
    
    return sections
